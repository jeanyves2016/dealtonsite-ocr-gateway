from PIL import Image
from fastapi import FastAPI, UploadFile, File
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pdfminer.high_level import extract_text
import pandas as pd
from docx import Document
import shutil
import subprocess
import os
import uuid
import re
import json

# 🔥 DB
from sqlalchemy import create_engine, Column, Integer, String, Text
from sqlalchemy.orm import declarative_base, sessionmaker

app = FastAPI(title="Dealtonsite OCR Gateway")

OUTPUT_DIR = "/app/output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# =====================
# 🟣 DATABASE
# =====================
DATABASE_URL = "sqlite:///./ocr.db"

engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(bind=engine)
Base = declarative_base()


class DocumentDB(Base):
    __tablename__ = "documents"

    id = Column(Integer, primary_key=True, index=True)
    filename = Column(String)
    text = Column(Text)
    fields = Column(Text)


Base.metadata.create_all(bind=engine)


@app.get("/health")
def health():
    return {"status": "ok"}


# =====================
# 🟣 UTILS
# =====================
def clean_text(text):
    return re.sub(r"[\x00-\x1F\x7F]", "", text)


def extract_fields(text):
    data = {}

    invoice = re.search(r"(facture|invoice)[^\n]*?(\d+)", text, re.IGNORECASE)
    if invoice:
        data["invoice_number"] = invoice.group(2)

    date = re.search(r"\b\d{2}/\d{2}/\d{4}\b", text)
    if date:
        data["date"] = date.group()

    amount = re.search(r"(\d+[.,]?\d*)\s?(FCFA|€|\$)", text)
    if amount:
        data["total_amount"] = amount.group(1)
        data["currency"] = amount.group(2)

    return data


# =====================
# 🟣 OCR
# =====================
@app.post("/ocr")
async def ocr(file: UploadFile = File(...)):
    file_id = str(uuid.uuid4())

    input_path = f"/tmp/{file_id}_{file.filename}"
    output_path = f"{OUTPUT_DIR}/ocr_{file_id}.pdf"

    # sauvegarde
    with open(input_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # conversion image → PDF
    if input_path.lower().endswith((".png", ".jpg", ".jpeg")):
        image = Image.open(input_path).convert("RGB")
        pdf_input_path = input_path + ".pdf"
        image.save(pdf_input_path)
        input_path = pdf_input_path

    # OCR
    try:
        subprocess.run([
            "ocrmypdf",
            "--force-ocr",
            input_path,
            output_path
        ], check=True)
    except subprocess.CalledProcessError as e:
        print("OCR ERROR:", e)
        return {"status": "error", "message": "OCR failed"}

    # extraction texte
    try:
        extracted_text = extract_text(output_path)
    except Exception as e:
        print("TEXT ERROR:", e)
        extracted_text = "Extraction texte échouée"

    cleaned_text = clean_text(extracted_text)
    fields = extract_fields(cleaned_text)

    # =====================
    # 🟣 SAVE TO DB
    # =====================
    db = SessionLocal()
    doc = DocumentDB(
        filename=file.filename,
        text=cleaned_text,
        fields=json.dumps(fields)
    )
    db.add(doc)
    db.commit()
    db.close()

    # Excel
    excel_path = f"{OUTPUT_DIR}/ocr_{file_id}.xlsx"
    df = pd.DataFrame([fields if fields else {"text": cleaned_text[:2000]}])
    df.to_excel(excel_path, index=False)

    # Word
    word_path = f"{OUTPUT_DIR}/ocr_{file_id}.docx"
    docx_file = Document()
    docx_file.add_heading("OCR Result", 0)
    docx_file.add_paragraph(cleaned_text)
    docx_file.save(word_path)

    return {
        "status": "success",
        "fields": fields,
        "text_preview": cleaned_text[:500],
        "download_pdf": f"/download/{file_id}",
        "download_excel": f"/download/excel/{file_id}",
        "download_word": f"/download/word/{file_id}"
    }


# =====================
# 🟣 DOWNLOAD
# =====================
@app.get("/download/{file_id}")
def download_pdf(file_id: str):
    file_path = f"{OUTPUT_DIR}/ocr_{file_id}.pdf"
    return FileResponse(file_path, media_type='application/pdf', filename="ocr_result.pdf")


@app.get("/download/excel/{file_id}")
def download_excel(file_id: str):
    file_path = f"{OUTPUT_DIR}/ocr_{file_id}.xlsx"
    return FileResponse(file_path, filename="result.xlsx")


@app.get("/download/word/{file_id}")
def download_word(file_id: str):
    file_path = f"{OUTPUT_DIR}/ocr_{file_id}.docx"
    return FileResponse(file_path, filename="result.docx")


# =====================
# 🟣 HISTORY
# =====================
@app.get("/documents")
def get_documents():
    db = SessionLocal()
    docs = db.query(DocumentDB).all()
    db.close()

    return [
        {
            "id": d.id,
            "filename": d.filename,
            "fields": json.loads(d.fields) if d.fields else {}
        }
        for d in docs
    ]


# =====================
# 🟣 FRONTEND
# =====================
app.mount("/", StaticFiles(directory="app/static", html=True), name="static")